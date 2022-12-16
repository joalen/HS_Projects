"""
Creates a reference document that maps one column for images while the other column designates the chapter and page associated with a PDF.

Features available: 
 - Extract images from a PDF 
 - Custom configurations (PDF path & Directory for extracted images) using a YAML file 
 - Write to Microsoft Word Document or Google Docs
 - Console-based GUI
"""

__author__ = "Alen Jo" 
__copyright__ = "Copyright 2022, Alen Jo" 

import fitz
import os.path 
from progress.bar import IncrementalBar
import time 
from ruamel.yaml import YAML

# edit these properties if you do NOT want to use a YAML file
pdfLocation = ""
imagesDir = ""
yaml = YAML() 

# Checks to see if the given configuration file is present. Do note that you will need to edit some properties in the given file
if os.path.exists('configs.yaml'):
    with open("configs.yaml", "r") as ymlfile: 
        configs = yaml.load(ymlfile) 
        pdfLocation = str(configs['configurations']['filepath'])
        imagesDir = str(configs['configurations']['compileDestination'])

pdf_file = fitz.open(pdfLocation)

def extractImages(start=0, end=len(pdf_file)):  
    """
    Function extracts all images from a PDF file, and then removes all images that have less than 6500 colors present in the image (removes symbols or redundant photos). 
    
    :param start: The page number to start extracting images from, defaults to 0 (optional)
    :param end: The last page of the PDF to extract images from
    """

    import os 
    bar = IncrementalBar('Extracting Images', max=((end-1) - (start-1)), suffix='%(percent).1f%% - %(eta)ds')   
    imageNumber = 0
    for page_index in range(start - 1, end - 1):
        for img in pdf_file.get_page_images(page_index):

            xref = img[0]
            smask = img[1] 

            if smask > 0:
                pix0 = fitz.Pixmap(pdf_file.extract_image(xref)["image"])
                if pix0.alpha:  # catch irregular situation
                    pix0 = fitz.Pixmap(pix0, 0)  # remove alpha channel
                mask = fitz.Pixmap(pdf_file.extract_image(smask)["image"])

                try:
                    pix = fitz.Pixmap(pix0, mask)
                except:  # fallback to original base image in case of problems
                    pix = fitz.Pixmap(pdf_file.extract_image(xref)["image"])

                ext = "png"


                imgfile = os.path.join(imagesDir, "%i.%s" % (imageNumber, ext))
                fout = open(imgfile, "wb")  

                try: 
                    fout.write(pix.tobytes(ext))
                except: 
                    continue 
                fout.close()
                
            if "/ColorSpace" in pdf_file.xref_object(xref, compressed=True):
                pix = fitz.Pixmap(pdf_file, xref)
                pix = fitz.Pixmap(fitz.csRGB, pix)        

                imgfile = os.path.join(imagesDir, "%i.%s" % (imageNumber, "png"))
                fout = open(imgfile, "wb")
                fout.write(pix.tobytes("png"))
                fout.close()
        
            bar.next()
            imageNumber += 1
    bar.finish() 

    from PIL import Image
    from collections import Counter

    for image in os.listdir(imagesDir):
        img = Image.open(imagesDir+image)
        colors = Counter(img.getdata()) 
        if len(colors) < 6500:
            os.remove(imagesDir+image)


def writeToWordDocument(title, start=0, end=len(pdf_file), chapter_number=None, photoRepetition=None):
    """
    It takes a title, start and end page numbers, chapter number, and a list of how many times each page
    should be repeated. It then creates a table with two columns and two rows, and then adds a row for
    each page in the range, and adds the page number to the second column. This function (in the name) 
    uses its development to create a Word Document instead of a Google Document. 
    
    :param title: The title of the document
    :param start: the starting page of the chapter, defaults to 0 (optional)
    :param end: the last page of the chapter
    :param chapter_number: The chapter number of the book
    :param photoRepetition: This is a list of integers that represent how many times each photo should
    be repeated
    """

    from docx import Document

    current_imageRepeatsIndex = 0
    document = Document()
    a = document.add_table(2, 2) 
    a.style = 'Table Grid'

    imageRepeats = photoRepetition

    rowIter = 0
    for i in range(start, end + 1):
        for j in range(imageRepeats[current_imageRepeatsIndex]):
            a.add_row()
            a.cell(rowIter, 2).add_paragraph()
            a.cell(rowIter, 3).add_paragraph(f"Chapter {chapter_number} - pg. {i}")
            rowIter += 1
            
        current_imageRepeatsIndex += 1
    
    a._tbl.remove(a.rows[0]._tr)
    a._tbl.remove(a.rows[1]._tr)
    document.save(f'{title}.docx')


def writeToGoogleDocs(start=0, end=len(pdf_file), chapter_number=None, photoRepetition=None, title=None, docID=None):
    from googleapiclient.discovery import build
    from gdoctableapppy import gdoctableapp
    from gdocsOperations import checkGoogleOAuth
    
    creds = None
    SCOPES = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/drive.file']
    # The file token.json stores the user's access and refresh tokens, and is
    # created automatically when the authorization flow completes for the first
    # time.
    creds = checkGoogleOAuth(SCOPES)

    service = build('docs', 'v1', credentials=creds)
    uniqueIdentifiers = []


    if docID == None:
        from gdocsOperations import createTable
        body = {'title': title}
        doc = service.documents().create(body=body).execute()

        DOCUMENTID = doc['documentId']
        createTable(creds, DOCUMENTID, imageRepeats=photoRepetition, strt=start, end_=end, chapterNumber=chapter_number, imagesDirectory=imagesDir)
    
    else:
        from gdocsOperations import appendTable
        appendTable(creds, DOCUMENTID, imageRepeats=photoRepetition, strt=start, end_=end, chapterNumber=chapter_number, imagesDirectory=imagesDir)

    def quickReplaceImage(path): 
        """
        The function takes a path to a directory of images, and replaces the placeholder text in the Google Doc with the images in the directory. 
        
        :param path: the path to the directory containing the images you want to replace
        """
        for i in range(len(os.listdir(path))):
            fileObjectList = [os.path.splitext(filename)[0] for filename in os.listdir(imagesDir)]
            newPath = imagesDir + sorted(fileObjectList, key = int)[i] + ".png"
            resource = { "oauth2": creds,"documentId": docID,"showAPIResponse": True,  "tableOnly": True, "searchText": f"--sample-- {uniqueIdentifiers[i]}", "replaceImageFilePath": newPath}
            gdoctableapp.ReplaceTextsToImages(resource)
            bar_inst_1.next() 
            time.sleep(1) 

    bar_inst_1= IncrementalBar('Adding Info  ', max=(end + 1 - start + 1), suffix='%(percent).1f%% - %(eta)ds')  
    quickReplaceImage(imagesDir)
    bar_inst_1.finish() 

    if title == None:
        print("Compiled successfully!")
    else:
        print(f"Compiled {title} successfully!")
    
    time.sleep(2)

def editConfigs(): 
    """
    If the configs.yaml file doesn't exist, create it and ask the user for the filepath and destination.
    
    If it does exist, update configurations by loading the filepath and destination from the configs.yaml file.
    """

    if not os.path.exists("configs.yaml"): 
        filePath = input("4/Insert PDF Path> ")
        compileDestination  = input("4/Picture Destination> ")
        
        configurations = """
        configurations: 
            filepath: {path}
            compileDestination: {dest}
        """.format(path=filePath, dest=compileDestination)
        
        with open("configs.yaml", "w") as ymlfile: 
            yaml.dump(yaml.load(configurations), ymlfile)

        editConfigs() 

    else: 
        with open("configs.yaml", "r") as ymlfile: 
            configs = yaml.load(ymlfile) 
        global file 
        global imagesDir
        file = str(configs['configurations']['filepath'])
        imagesDir = str(configs['configurations']['compileDestination'])
    
    

if __name__ == "__main__":
    import os
    print("Two-Column Image-To-Text Reference Page Generator")
    print("")
    
    time.sleep(3)
    os.system('cls')

    def menu(): 
        menuList = ["Menu:", "1. Extract Images", "2. Write Content (using Google Docs)", "3. Write Content (using Microsoft Word)", "4. Edit Configurations", "5. Exit"]
        print(*menuList, sep="\n")

    menu()
    c = int(input("> "))

    while (c != 5):
        match c:
            case 1: 
                startPg = int(input("1/Start Page> ")) 
                endPg = int(input("1/End Page> "))
                extractImages(startPg, (endPg + 1))
                os.system('cls')
                menu()
                c = int(input("> "))
            case 2: 
                repeat = input("Repeat? (y/n)")
                firstTime = False
                while repeat != "n":
                    if firstTime == True:
                        startPg = int(input("2/Start Page> ")) 
                        endPg = int(input("2/End Page> "))
                        repeatingPhotos = input("2/Photo Repeats> ").split()
                        existDocID = input("2/Insert Doc ID> ")
                        chapterNumber = int(input("2/Chapter Number> "))
                        writeToGoogleDocs(startPg, endPg, chapterNumber, [eval(i) for i in repeatingPhotos], None, existDocID)
                    else:
                        startPg = int(input("2/Start Page> ")) 
                        endPg = int(input("2/End Page> "))
                        repeatingPhotos = input("2/Photo Repeats> ").split()
                        titleForDoc = input("2/Title of Doc> ")
                        chapterNumber = int(input("2/Chapter Number> "))
                        writeToGoogleDocs(startPg, endPg, chapterNumber, [eval(i) for i in repeatingPhotos], titleForDoc)
                    
                    os.system('cls')
                    repeat = input("Repeat? (y/n)")
                    firstTime = True
                os.system('cls')
                menu()
                c = int(input("> "))
            case 3: 
                startPg = int(input("3/Start Page> ")) 
                endPg = int(input("3/End Page> "))
                repeatingPhotos = input("3/Photo Repeats> ").split()
                docName = input("3/Name of Document> ")
                writeToWordDocument(docName, startPg, endPg, 15, [eval(i) for i in repeatingPhotos])
                os.system('cls')
                menu()
                c = int(input("> "))       
            case 4: 
                editConfigs()
                os.system('cls')
                menu() 
                c = int(input("> "))
            case _:
                print("Try again!")
                os.system('cls')
                menu()
                c = int(input("> "))
